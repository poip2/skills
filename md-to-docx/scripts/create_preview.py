#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(SCRIPT_DIR, 'template.docx')
OUTPUT_PATH = os.path.join(SCRIPT_DIR, 'template_preview.docx')

def set_run_font(run, font_cn='宋体', font_en='Times New Roman', size=12, bold=False, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = font_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)

def add_shading(paragraph, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    paragraph._p.get_or_add_pPr().append(shading)

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_preview():
    doc = Document(TEMPLATE_PATH)
    
    doc.add_heading('格式规范预览文档', level=1)
    p = doc.add_paragraph()
    run = p.add_run('本文档展示 md-to-docx Skill 支持的所有格式规范。')
    set_run_font(run)
    
    doc.add_heading('一、标题样式', level=1)
    
    h1 = doc.add_heading('一级标题 (Heading 1)', level=1)
    p = doc.add_paragraph()
    run = p.add_run('字号：22pt（二号），加粗，段前自动分页')
    set_run_font(run, size=10, italic=True)
    
    h2 = doc.add_heading('二级标题 (Heading 2)', level=2)
    p = doc.add_paragraph()
    run = p.add_run('字号：16pt（三号），加粗，不分页')
    set_run_font(run, size=10, italic=True)
    
    h3 = doc.add_heading('三级标题 (Heading 3)', level=3)
    p = doc.add_paragraph()
    run = p.add_run('字号：15pt（小三），加粗，不分页')
    set_run_font(run, size=10, italic=True)
    
    h4 = doc.add_heading('四级标题 (Heading 4)', level=4)
    p = doc.add_paragraph()
    run = p.add_run('字号：14pt（四号），加粗，不分页')
    set_run_font(run, size=10, italic=True)
    
    h5 = doc.add_heading('五级标题 (Heading 5)', level=5)
    p = doc.add_paragraph()
    run = p.add_run('字号：14pt（四号），加粗，不分页')
    set_run_font(run, size=10, italic=True)
    
    h6 = doc.add_heading('六级标题 (Heading 6)', level=6)
    p = doc.add_paragraph()
    run = p.add_run('字号：12pt（小四），加粗，不分页')
    set_run_font(run, size=10, italic=True)
    
    doc.add_heading('二、正文样式', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('正文段落：宋体 + Times New Roman，12pt（小四），首行缩进 0.74cm，1.5 倍行距。')
    set_run_font(run)
    
    p = doc.add_paragraph()
    run = p.add_run('粗体文本：')
    set_run_font(run)
    run = p.add_run('这是粗体文本')
    set_run_font(run, bold=True)
    
    p = doc.add_paragraph()
    run = p.add_run('斜体文本：')
    set_run_font(run)
    run = p.add_run('这是斜体文本')
    set_run_font(run, italic=True)
    
    p = doc.add_paragraph()
    run = p.add_run('粗体+斜体：')
    set_run_font(run)
    run = p.add_run('这是粗体斜体文本')
    set_run_font(run, bold=True, italic=True)
    
    doc.add_heading('三、列表样式', level=1)
    
    doc.add_heading('3.1 无序列表', level=2)
    for item in ['无序列表项 1', '无序列表项 2', '无序列表项 3']:
        p = doc.add_paragraph()
        run = p.add_run(f'• {item}')
        set_run_font(run)
        p.paragraph_format.left_indent = Cm(0.74)
    
    doc.add_heading('3.2 有序列表', level=2)
    for i, item in enumerate(['有序列表项 1', '有序列表项 2', '有序列表项 3'], 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {item}')
        set_run_font(run)
        p.paragraph_format.left_indent = Cm(0.74)
    
    doc.add_heading('3.3 多级列表', level=2)
    p = doc.add_paragraph()
    run = p.add_run('• 一级列表项')
    set_run_font(run)
    p.paragraph_format.left_indent = Cm(0.74)
    
    p = doc.add_paragraph()
    run = p.add_run('• 二级列表项')
    set_run_font(run)
    p.paragraph_format.left_indent = Cm(1.48)
    
    p = doc.add_paragraph()
    run = p.add_run('• 三级列表项')
    set_run_font(run)
    p.paragraph_format.left_indent = Cm(2.22)
    
    doc.add_heading('四、表格样式', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('表格样式：Table Grid，居中显示，表头背景 #D9D9D9（浅灰色），表头居中，数据左对齐。')
    set_run_font(run, size=10, italic=True)
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['列1', '列2', '列3']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
        set_cell_shading(cell, 'D9D9D9')
    
    data = [
        ['数据1', '数据2', '数据3'],
        ['数据4', '数据5', '数据6'],
        ['数据7', '数据8', '数据9'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_text
    
    doc.add_heading('五、代码块样式', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('代码块：Consolas 字体，9pt（小五），背景色 #F5F5F5，左缩进 0.5cm')
    set_run_font(run, size=10, italic=True)
    
    p = doc.add_paragraph()
    run = p.add_run('[python]')
    set_run_font(run, font_en='Consolas', size=9, italic=True)
    
    code_lines = [
        'def hello_world():',
        '    print("Hello, World!")',
        '    return True',
    ]
    for line in code_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_run_font(run, font_cn='Consolas', font_en='Consolas', size=9)
        p.paragraph_format.left_indent = Cm(0.5)
        add_shading(p, 'F5F5F5')
    
    doc.add_heading('六、行内代码样式', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('行内代码：')
    set_run_font(run)
    run = p.add_run('print("Hello")')
    set_run_font(run, font_cn='Consolas', font_en='Consolas', size=12)
    
    p = doc.add_paragraph()
    run = p.add_run('行内代码：Consolas 字体，12pt（小四），背景色 #F0F0F0')
    set_run_font(run, size=10, italic=True)
    
    doc.add_heading('七、引用块样式', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('引用块：左边框 #6366F1（紫色），边框宽度 1.5pt，左右缩进 1cm，斜体')
    set_run_font(run, size=10, italic=True)
    
    p = doc.add_paragraph()
    run = p.add_run('这是一段引用文字，使用斜体样式，左边有紫色竖线标识。')
    set_run_font(run, italic=True)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    
    doc.add_heading('八、分隔线样式', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('分隔线：底部边框样式，颜色 #CCCCCC（浅灰色），段前段后间距 6pt')
    set_run_font(run, size=10, italic=True)
    
    p = doc.add_paragraph()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph('分隔线上方内容')
    doc.add_paragraph('分隔线下方内容')
    
    doc.add_heading('九、字体规范汇总', level=1)
    
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['元素类型', '中文字体', '英文字体', '字号']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
        set_cell_shading(cell, 'D9D9D9')
    
    data = [
        ['正文', '宋体', 'Times New Roman', '12pt（小四）'],
        ['一级标题', '宋体', 'Times New Roman', '22pt（二号）'],
        ['二级标题', '宋体', 'Times New Roman', '16pt（三号）'],
        ['三级标题', '宋体', 'Times New Roman', '15pt（小三）'],
        ['四级标题', '宋体', 'Times New Roman', '14pt（四号）'],
        ['代码块', 'Consolas', 'Consolas', '9pt（小五）'],
        ['行内代码', 'Consolas', 'Consolas', '12pt（小四）'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_text
    
    doc.add_heading('十、段落规范汇总', level=1)
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['属性', '设置值', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
        set_cell_shading(cell, 'D9D9D9')
    
    data = [
        ['首行缩进', '0.74cm', '约两个汉字宽度'],
        ['行间距', '1.5 倍', '提升阅读舒适度'],
        ['段前间距', '0pt', '保持紧凑排版'],
        ['段后间距', '0pt', '保持紧凑排版'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_text
    
    doc.add_heading('十一、分页控制', level=1)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['规则', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
        set_cell_shading(cell, 'D9D9D9')
    
    data = [
        ['一级标题前分页', '每个一级标题自动另起一页'],
        ['其他标题不分页', '二级及以下标题保持连续'],
        ['封面页后分页', '封面页结束后自动分页'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_text
    
    doc.save(OUTPUT_PATH)
    print(f'完整格式规范预览文档已创建: {OUTPUT_PATH}')

if __name__ == '__main__':
    create_preview()
